"""
Microsoft Office (Excel/Word) control tools.

Uses openpyxl for Excel and python-docx for Word.
"""

from langchain_core.tools import tool
from pathlib import Path
from typing import Optional
import csv
import io

# Limits
MAX_ROWS_DISPLAY = 100
MAX_COLS_DISPLAY = 26  # A-Z


def _resolve_path(file_path: str) -> Path:
    """Resolve and validate file path."""
    return Path(file_path).expanduser().resolve()


def _format_table(rows: list[list[str]], headers: Optional[list[str]] = None) -> str:
    """Format data as a markdown table."""
    if not rows:
        return "(empty)"

    # Use first row as headers if not provided
    if headers is None and rows:
        headers = [f"Col{i+1}" for i in range(len(rows[0]))]

    # Calculate column widths
    all_rows = [headers] + rows if headers else rows
    col_widths = []
    for col_idx in range(len(all_rows[0])):
        width = max(len(str(row[col_idx])) for row in all_rows if col_idx < len(row))
        col_widths.append(min(width, 50))  # Max 50 chars per column

    # Format header
    output = []
    if headers:
        header_line = "| " + " | ".join(
            str(h)[:50].ljust(col_widths[i]) for i, h in enumerate(headers)
        ) + " |"
        separator = "|-" + "-|-".join("-" * w for w in col_widths) + "-|"
        output.extend([header_line, separator])

    # Format rows
    for row in rows:
        row_line = "| " + " | ".join(
            str(row[i] if i < len(row) else "")[:50].ljust(col_widths[i])
            for i in range(len(col_widths))
        ) + " |"
        output.append(row_line)

    return "\n".join(output)


@tool
def read_excel(file_path: str, sheet_name: Optional[str] = None) -> str:
    """
    Read an Excel file and return its contents as a formatted table.

    Args:
        file_path: Path to the Excel file (.xlsx, .xls)
        sheet_name: Name of the sheet to read (default: first sheet)

    Returns:
        Contents as a markdown table

    Examples:
        - read_excel("C:\\Users\\user\\Documents\\report.xlsx")
        - read_excel("~/data.xlsx", sheet_name="Sheet2")

    Limits:
        - Maximum 100 rows displayed
        - Maximum 26 columns (A-Z)
    """
    try:
        from openpyxl import load_workbook

        path = _resolve_path(file_path)

        if not path.exists():
            return f"Error: File not found: {path}"

        if not path.suffix.lower() in [".xlsx", ".xlsm", ".xltx", ".xltm"]:
            return f"Error: Not an Excel file (.xlsx): {path.suffix}"

        # Load workbook
        wb = load_workbook(path, read_only=True, data_only=True)

        # Get sheet
        if sheet_name:
            if sheet_name not in wb.sheetnames:
                return f"Error: Sheet '{sheet_name}' not found. Available: {', '.join(wb.sheetnames)}"
            ws = wb[sheet_name]
        else:
            ws = wb.active
            sheet_name = ws.title

        # Read data
        rows = []
        headers = None

        for row_idx, row in enumerate(ws.iter_rows(max_row=MAX_ROWS_DISPLAY + 1, max_col=MAX_COLS_DISPLAY)):
            row_data = [str(cell.value) if cell.value is not None else "" for cell in row]

            # Skip completely empty rows
            if not any(row_data):
                continue

            if row_idx == 0:
                headers = row_data
            else:
                rows.append(row_data)

        wb.close()

        if not rows and not headers:
            return f"Sheet '{sheet_name}' is empty"

        # Format output
        total_rows = ws.max_row or 0
        total_cols = ws.max_column or 0

        output = [f"File: {path.name}"]
        output.append(f"Sheet: {sheet_name}")
        output.append(f"Size: {total_rows} rows x {total_cols} columns")

        if total_rows > MAX_ROWS_DISPLAY:
            output.append(f"(showing first {MAX_ROWS_DISPLAY} rows)")

        output.append("")
        output.append(_format_table(rows, headers))

        return "\n".join(output)

    except ImportError:
        return "Error: 'openpyxl' not installed. Run: pip install openpyxl"
    except Exception as e:
        return f"Error reading Excel: {type(e).__name__}: {str(e)}"


@tool
def write_excel(file_path: str, data: str, sheet_name: str = "Sheet1") -> str:
    """
    Write data to an Excel file.

    Args:
        file_path: Path to save the Excel file (.xlsx)
        data: CSV-formatted data (comma-separated, newline for rows)
        sheet_name: Name for the sheet (default: Sheet1)

    Returns:
        Success message with file path

    Examples:
        - write_excel("output.xlsx", "Name,Age,City\\nAlice,30,Seoul\\nBob,25,Busan")
        - write_excel("~/report.xlsx", "A,B,C\\n1,2,3\\n4,5,6", sheet_name="Data")

    Note:
        First row is treated as headers
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font

        path = _resolve_path(file_path)

        # Ensure .xlsx extension
        if not path.suffix.lower() == ".xlsx":
            path = path.with_suffix(".xlsx")

        # Parse CSV data
        reader = csv.reader(io.StringIO(data))
        rows = list(reader)

        if not rows:
            return "Error: No data provided"

        # Create workbook
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name

        # Write data
        for row_idx, row in enumerate(rows, 1):
            for col_idx, value in enumerate(row, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                # Bold first row (headers)
                if row_idx == 1:
                    cell.font = Font(bold=True)

        # Auto-adjust column widths
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            ws.column_dimensions[column].width = min(max_length + 2, 50)

        # Ensure parent directory exists
        path.parent.mkdir(parents=True, exist_ok=True)

        # Save
        wb.save(path)
        wb.close()

        return f"Successfully saved Excel file: {path}\nRows: {len(rows)}, Columns: {len(rows[0]) if rows else 0}"

    except ImportError:
        return "Error: 'openpyxl' not installed. Run: pip install openpyxl"
    except Exception as e:
        return f"Error writing Excel: {type(e).__name__}: {str(e)}"


@tool
def read_word(file_path: str) -> str:
    """
    Read a Word document (.docx) and return its text content.

    Args:
        file_path: Path to the Word document

    Returns:
        Document text content

    Examples:
        - read_word("C:\\Users\\user\\Documents\\report.docx")
        - read_word("~/letter.docx")
    """
    try:
        from docx import Document

        path = _resolve_path(file_path)

        if not path.exists():
            return f"Error: File not found: {path}"

        if not path.suffix.lower() == ".docx":
            return f"Error: Not a Word document (.docx): {path.suffix}"

        # Load document
        doc = Document(path)

        # Extract text
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract tables
        tables_text = []
        for table_idx, table in enumerate(doc.tables):
            table_rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_rows.append(row_data)
            if table_rows:
                tables_text.append(f"\n[Table {table_idx + 1}]\n{_format_table(table_rows[1:], table_rows[0] if table_rows else None)}")

        # Combine output
        output = [f"File: {path.name}"]
        output.append(f"Paragraphs: {len(paragraphs)}, Tables: {len(doc.tables)}")
        output.append("")
        output.append("\n\n".join(paragraphs))

        if tables_text:
            output.append("\n" + "\n".join(tables_text))

        return "\n".join(output)

    except ImportError:
        return "Error: 'python-docx' not installed. Run: pip install python-docx"
    except Exception as e:
        return f"Error reading Word document: {type(e).__name__}: {str(e)}"


@tool
def write_word(file_path: str, content: str, title: Optional[str] = None) -> str:
    """
    Create a Word document (.docx) with the given content.

    Args:
        file_path: Path to save the Word document
        content: Text content (paragraphs separated by blank lines)
        title: Optional document title (will be added as heading)

    Returns:
        Success message with file path

    Examples:
        - write_word("report.docx", "This is the first paragraph.\\n\\nThis is the second paragraph.")
        - write_word("~/memo.docx", "Meeting notes here...", title="Meeting Notes")
    """
    try:
        from docx import Document
        from docx.shared import Pt

        path = _resolve_path(file_path)

        # Ensure .docx extension
        if not path.suffix.lower() == ".docx":
            path = path.with_suffix(".docx")

        # Create document
        doc = Document()

        # Add title
        if title:
            heading = doc.add_heading(title, level=0)

        # Add paragraphs
        paragraphs = content.split("\n\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                # Check if it looks like a heading (starts with # markdown style)
                if para_text.startswith("# "):
                    doc.add_heading(para_text[2:], level=1)
                elif para_text.startswith("## "):
                    doc.add_heading(para_text[3:], level=2)
                elif para_text.startswith("### "):
                    doc.add_heading(para_text[4:], level=3)
                else:
                    doc.add_paragraph(para_text)

        # Ensure parent directory exists
        path.parent.mkdir(parents=True, exist_ok=True)

        # Save
        doc.save(path)

        return f"Successfully saved Word document: {path}"

    except ImportError:
        return "Error: 'python-docx' not installed. Run: pip install python-docx"
    except Exception as e:
        return f"Error writing Word document: {type(e).__name__}: {str(e)}"


@tool
def list_excel_sheets(file_path: str) -> str:
    """
    List all sheets in an Excel file.

    Args:
        file_path: Path to the Excel file

    Returns:
        List of sheet names
    """
    try:
        from openpyxl import load_workbook

        path = _resolve_path(file_path)

        if not path.exists():
            return f"Error: File not found: {path}"

        wb = load_workbook(path, read_only=True)
        sheets = wb.sheetnames
        wb.close()

        output = [f"File: {path.name}"]
        output.append(f"Sheets ({len(sheets)}):")
        for i, name in enumerate(sheets, 1):
            output.append(f"  {i}. {name}")

        return "\n".join(output)

    except ImportError:
        return "Error: 'openpyxl' not installed. Run: pip install openpyxl"
    except Exception as e:
        return f"Error: {type(e).__name__}: {str(e)}"
